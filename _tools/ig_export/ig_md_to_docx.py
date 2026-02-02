#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE


FNAME_RE = re.compile(r"^(?P<date>\d{4}-\d{2}-\d{2})(?:__(?P<ord>\d+))?\.md$")


@dataclass(frozen=True)
class PostFile:
    path: Path
    date: str  # YYYY-MM-DD
    ordinal: int  # 1..N
    year: str  # YYYY


def parse_post_file(path: Path) -> Optional[PostFile]:
    m = FNAME_RE.match(path.name)
    if not m:
        return None
    date = m.group("date")
    ordinal = int(m.group("ord") or "1")
    year = date[:4]
    return PostFile(path=path, date=date, ordinal=ordinal, year=year)


def read_url_and_text(md_path: Path) -> tuple[str, str]:
    """
    Our md format:
      line1: URL
      blank
      rest: caption text
    """
    raw = md_path.read_text(encoding="utf-8")
    lines = raw.splitlines()
    url = (lines[0].strip() if lines else "").strip()
    text = "\n".join(lines[2:]).strip() if len(lines) >= 3 else ""
    return url, text


def add_hyperlink(paragraph, url: str, text: Optional[str] = None) -> None:
    """
    Add a real clickable hyperlink to a paragraph.
    python-docx doesn't expose a public API for this, so we add the underlying XML.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    # Blue + underlined (typical hyperlink styling)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    r_pr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text or url
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def main() -> int:
    ap = argparse.ArgumentParser(description="Combine IG caption .md files into a single .docx grouped by year.")
    ap.add_argument("--md-dir", required=True, help="Directory with per-post .md files (YYYY-MM-DD.md, YYYY-MM-DD__2.md...).")
    ap.add_argument("--out-docx", required=True, help="Output .docx path.")
    args = ap.parse_args()

    md_dir = Path(args.md_dir)
    out_docx = Path(args.out_docx)
    out_docx.parent.mkdir(parents=True, exist_ok=True)

    post_files: list[PostFile] = []
    for p in md_dir.iterdir():
        if not p.is_file():
            continue
        pf = parse_post_file(p)
        if pf:
            post_files.append(pf)

    post_files.sort(key=lambda x: (x.year, x.date, x.ordinal))

    doc = Document()
    doc.add_heading("Instagram posts (captions)", level=0)

    current_year: Optional[str] = None
    for pf in post_files:
        if pf.year != current_year:
            current_year = pf.year
            doc.add_page_break()
            doc.add_heading(current_year, level=1)

        url, text = read_url_and_text(pf.path)

        # Date line
        title = pf.date if pf.ordinal == 1 else f"{pf.date} ({pf.ordinal})"
        doc.add_heading(title, level=2)

        # URL line
        if url:
            p = doc.add_paragraph()
            add_hyperlink(p, url)

        # Text only (split by blank lines -> paragraphs)
        if text:
            chunks = [c.strip() for c in re.split(r"\n\s*\n", text) if c.strip()]
            for c in chunks:
                doc.add_paragraph(c)
        else:
            doc.add_paragraph("")

    doc.save(str(out_docx))
    print(f"[done] wrote {out_docx} from {len(post_files)} posts")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())


